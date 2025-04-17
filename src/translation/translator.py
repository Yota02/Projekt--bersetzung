from pdf2docx import Converter
from docx import Document
from deep_translator import GoogleTranslator
import time
import os

class PDFTranslator:
    def __init__(self, callback=None):
        self.callback = callback
        
    def translate(self, pdf_path, output_path, target_lang='fr'):
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Le fichier {pdf_path} n'existe pas.")

        temp_docx = 'temp_document.docx'
        
        self._log("Conversion du PDF en DOCX...")
        cv = Converter(pdf_path)
        cv.convert(temp_docx)
        cv.close()
        
        self._log("Début de la traduction...")
        doc = Document(temp_docx)
        
        # Convertir les noms de langues
        lang_map = {
    'German': 'de',
    'Japanese': 'ja',
    'French': 'fr',
    'English': 'en',
    'Russian': 'ru',
    'Italian': 'it'
}
        
        lang_code = lang_map.get(target_lang, None)
        if not lang_code:
            raise ValueError(f"Langue cible '{target_lang}' non prise en charge.")

        self._log("Début de la traduction...")
        translator = GoogleTranslator(source='auto', target=lang_code)
        
        total_paragraphs = len(doc.paragraphs)
        current_text = []
        
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            
            # Si le texte se termine par un point, c'est probablement une fin de phrase
            if text:
                # Ajoute le texte au paragraphe en cours
                current_text.append(text)
                
                # Vérifie si c'est la fin d'un paragraphe logique
                if text.endswith(('.', '!', '?', ':', ';')) or i == total_paragraphs:
                    full_text = ' '.join(current_text)
                    
                    try:
                        # Découpe en chunks de 4000 caractères maximum
                        chunks = [full_text[i:i+4000] for i in range(0, len(full_text), 4000)]
                        translated_chunks = []
                        
                        for chunk in chunks:
                            try:
                                translated_chunk = translator.translate(chunk)
                                translated_chunks.append(translated_chunk)
                            except Exception as e:
                                self._log(f"Erreur de traduction pour un segment: {str(e)}")
                                translated_chunks.append(chunk)
                        
                        # Met à jour le texte du dernier paragraphe du groupe
                        para.text = ' '.join(translated_chunks)
                        # Efface les paragraphes précédents du groupe
                        for prev_para in doc.paragraphs[i-len(current_text):i-1]:
                            prev_para.text = ''
                            
                        current_text = []  # Réinitialise pour le prochain groupe
                        
                    except Exception as e:
                        self._log(f"Erreur lors de la traduction du paragraphe {i}: {str(e)}")
                        continue
                        
                    progress = (i / total_paragraphs) * 100
                    self._log(f"Progression: {i}/{total_paragraphs} paragraphes traités")
                    if self.callback:
                        self.callback(progress)
    
        self._log("Sauvegarde du document traduit...")
        doc.save(output_path)
        
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
        
        self._log(f"Traduction terminée! Fichier sauvegardé sous: {output_path}")
        
    def _log(self, message):
        if self.callback:
            self.callback(message=message)
        print(message)